"""
shared/file_reader.py — Shared document ingestion layer.

This is a copy of code/file_reader.py promoted to a shared utility so
every agent and the orchestrator entry point can import it without
depending on the legacy `code/` directory.

Supports:
  • .txt  — plain text (UTF-8, latin-1 fallback)
  • .pdf  — pdfplumber (primary), pypdf (fallback)
  • .docx — python-docx (paragraphs + tables)
  • .doc  — mammoth conversion (no LibreOffice required)

Usage:
    from shared.file_reader import read_document, get_document_metadata
    text = read_document("paper.pdf")
"""

import logging
from pathlib import Path

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx", ".doc"}


def read_document(file_path: str) -> str:
    """
    Reads any supported document and returns its full text as a plain string.

    Args:
        file_path: Path to the file (.txt, .pdf, .docx, .doc)

    Returns:
        Extracted text as a string.

    Raises:
        ValueError: If the file type is not supported.
        FileNotFoundError: If the file does not exist.
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {path.resolve()}")

    suffix = path.suffix.lower()

    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type '{suffix}'. "
            f"Supported types: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    logger.info(f"Reading {suffix.upper()} file: {path.name}")

    if suffix == ".txt":
        return _read_txt(path)
    elif suffix == ".pdf":
        return _read_pdf(path)
    elif suffix in {".docx", ".doc"}:
        return _read_docx(path)


# ─── TXT ──────────────────────────────────────────────────────────────────────

def _read_txt(path: Path) -> str:
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        logger.warning("UTF-8 decode failed, retrying with latin-1")
        text = path.read_text(encoding="latin-1")
    logger.info(f"TXT: {len(text.split())} words extracted.")
    return text.strip()


# ─── PDF ──────────────────────────────────────────────────────────────────────

def _read_pdf(path: Path) -> str:
    text = ""

    # Primary: pdfplumber
    try:
        import pdfplumber

        with pdfplumber.open(str(path)) as pdf:
            pages_text = []
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    pages_text.append(page_text)
                else:
                    logger.debug(f"PDF page {i+1} yielded no text (image-only?).")
            text = "\n\n".join(pages_text).strip()

        if text:
            logger.info(f"PDF (pdfplumber): {len(pdf.pages)} pages, {len(text.split())} words.")
            return text

        logger.warning("pdfplumber returned empty text — trying pypdf fallback.")

    except ImportError:
        logger.warning("pdfplumber not installed — trying pypdf fallback.")
    except Exception as e:
        logger.warning(f"pdfplumber error: {e} — trying pypdf fallback.")

    # Fallback: pypdf
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        pages_text = [p.extract_text() for p in reader.pages if p.extract_text()]
        text = "\n\n".join(pages_text).strip()

        if text:
            logger.info(f"PDF (pypdf fallback): {len(reader.pages)} pages, {len(text.split())} words.")
            return text

        logger.warning("pypdf also returned empty text. PDF may be scanned/image-only.")
        return ""

    except ImportError:
        raise ImportError(
            "Neither pdfplumber nor pypdf is installed.\n"
            "Install with: pip install pdfplumber pypdf"
        )
    except Exception as e:
        logger.error(f"PDF extraction failed entirely: {e}")
        return ""


# ─── DOCX / DOC ───────────────────────────────────────────────────────────────

def _read_docx(path: Path) -> str:
    suffix = path.suffix.lower()

    if suffix == ".doc":
        return _read_doc_via_mammoth(path)

    try:
        from docx import Document

        doc = Document(str(path))
        sections: list[str] = []

        for para in doc.paragraphs:
            stripped = para.text.strip()
            if stripped:
                sections.append(stripped)

        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    sections.append(row_text)

        text = "\n".join(sections).strip()
        logger.info(f"DOCX: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables, {len(text.split())} words.")
        return text

    except ImportError:
        raise ImportError("python-docx is not installed. Install with: pip install python-docx")
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return ""


def _read_doc_via_mammoth(path: Path) -> str:
    try:
        import mammoth

        with open(str(path), "rb") as f:
            result = mammoth.extract_raw_text(f)

        text = result.value.strip()
        for msg in result.messages:
            logger.debug(f"mammoth: {msg}")

        logger.info(f"DOC (mammoth): {len(text.split())} words extracted.")
        return text

    except ImportError:
        raise ImportError("mammoth is not installed. Install with: pip install mammoth")
    except Exception as e:
        logger.error(f"DOC extraction via mammoth failed: {e}")
        return ""


# ─── Metadata ─────────────────────────────────────────────────────────────────

def get_document_metadata(file_path: str) -> dict:
    """Returns basic metadata about the document (size, type, pages)."""
    path = Path(file_path)
    suffix = path.suffix.lower()
    size_kb = path.stat().st_size / 1024

    meta = {
        "filename": path.name,
        "extension": suffix,
        "size_kb": round(size_kb, 2),
        "pages": None,
        "word_count": None,
    }

    if suffix == ".pdf":
        try:
            import pdfplumber
            with pdfplumber.open(str(path)) as pdf:
                meta["pages"] = len(pdf.pages)
        except Exception:
            pass

    elif suffix in {".docx", ".doc"}:
        meta["pages"] = "N/A"

    return meta
