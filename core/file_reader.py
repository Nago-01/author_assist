"""
Unified file ingestion layer — used by main.py CLI.
Supports PDF, DOCX, DOC, TXT.
Agents never call this directly — they always receive plain text from the pipeline.
"""

from __future__ import annotations

import os
from pathlib import Path


def read_file(path: str) -> str:
    """
    Read an article from disk and return its plain-text content.
    Raises ValueError for unsupported file types.
    Raises FileNotFoundError if the file does not exist.
    """
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"File not found: {path}")

    suffix = p.suffix.lower()

    if suffix == ".txt":
        return _read_txt(p)
    elif suffix == ".pdf":
        return _read_pdf(p)
    elif suffix == ".docx":
        return _read_docx(p)
    elif suffix == ".doc":
        return _read_doc(p)
    else:
        raise ValueError(
            f"Unsupported file type '{suffix}'. "
            "Supported: .txt, .pdf, .docx, .doc"
        )


# Readers

def _read_txt(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="latin-1")


def _read_pdf(path: Path) -> str:
    # Primary: pdfplumber
    try:
        import pdfplumber  

        text_parts: list[str] = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                extracted = page.extract_text()
                if extracted:
                    text_parts.append(extracted)
        text = "\n\n".join(text_parts)
        if text.strip():
            return text
    except ImportError:
        pass
    except Exception:  
        pass

    # Fallback: pypdf
    try:
        from pypdf import PdfReader  

        reader = PdfReader(str(path))
        pages = [page.extract_text() or "" for page in reader.pages]
        text = "\n\n".join(pages)
        if text.strip():
            return text
    except ImportError:
        pass

    return ""  


def _read_docx(path: Path) -> str:
    from docx import Document  

    doc = Document(str(path))
    parts: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)


def _read_doc(path: Path) -> str:
    import mammoth  

    with open(str(path), "rb") as f:
        result = mammoth.extract_raw_text(f)
    return result.value
